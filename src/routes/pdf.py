import uuid
import os
from typing import List
from fastapi import HTTPException
from fastapi import APIRouter, UploadFile, status, Request
from fastapi import Depends, Form, File, Header
from pikepdf import Pdf, Page, Encryption, PdfImage
from zipfile import ZipFile
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from bs4 import BeautifulSoup
from fastapi_limiter.depends import RateLimiter

from ..dependencies.jwtBearer import JWTBearer
from ..dependencies.packageLimit import PackageLimit
from ..dependencies.apiKeyAuthorization import ApiKeyAuthorization

from ..utils.validation import Validations, PdfPageSize
from ..utils.pdfTask import PdfTask, TaskError, parseFilesPassword
from ..utils.defers_collector import defers_collector, defer

from ..services.fileManagement import FileManagement, ALLOWED_MIME_TYPES
from ..services.imageHandling import ImageHandling
from ..services.pdfHandling import PdfHandling
from ..services.featureLimit import FeatureLimit
from ..services.pdfaltoHandling import PdfaltoHandling
from ..services.userManagement import UserManagement
from ..services.folderManagement import FolderManagement

from ..models.errorModel import ErrorInfoModel
from ..models.queryParams import pdfDeletePagesQueryParams, pdfRotateQueryParams, pdfSplitQueryParams, pdfSortQueryParams, pdfSignatureQueryParams
from ..models.platformFeature import APP_PLATFORM_FEATURE, WEB_PLATFORM_FEATURE
from ..models.planFeature import FREE_PLAN_FEATURE
from ..models.response import Response
# from ..services.backgroundJobs import BackgroundJobs
from ..config.config import settings



fMgm = FileManagement()
uMgm = UserManagement()
folderMgm = FolderManagement()
packageLimit = PackageLimit()
pdfRouter = APIRouter(
    dependencies=[
            Depends(ApiKeyAuthorization()),
            Depends(JWTBearer()), 
            Depends(RateLimiter(times=settings.RateLimitSettings.Times, seconds=settings.RateLimitSettings.Seconds)), 
            Depends(packageLimit.PdfManipulationLimitDependency)
        ]
)

@pdfRouter.post("/delete-pages", response_model=Response, dependencies=[Depends(packageLimit.FileLimitDependency)])
def pdfDeletePages(
    request: Request,
    query: pdfDeletePagesQueryParams = Depends(),
    file: UploadFile = File(...),
):
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")

    if len(query.pages) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="pages is required")

    pagesDeleteArr = query.pages.split('-')
    if len(pagesDeleteArr) != 1 and len(pagesDeleteArr) != 2:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="pages is failed")

    try:
        usr = request.state.userDict[u'email']
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan
        filePath = ImageHandling.saveImageToTmp(file)

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
            raise Exception("This feature is for premium account. Your current plan is free")
        
        # Delete pages file pdf
        with Pdf.open(filePath, allow_overwriting_input=True) as pdf:
            if len(pagesDeleteArr) == 1:
                listPageDelete = query.pages.split(',')
                for page in listPageDelete:
                    p = int(page) - 1
                    del pdf.pages[p]

            if len(pagesDeleteArr) == 2:
                t = int(pagesDeleteArr[0]) - 1
                f = int(pagesDeleteArr[1])
                del pdf.pages[t:f]

            pdf.save(filePath)

        # platform Web Add Watermark
        if platform == WEB_PLATFORM_FEATURE.name:
            filePath = PdfHandling().addWatermark(filePath)
        # Upload file in cloud
        fObj = fMgm.CreateFromPath(userId=usr, filePath=str(filePath))
        url = fMgm.GetDownloadLink(fObj=fObj)
        # Clear file 
        ImageHandling().cleanSingleImage(filePathWithName=str(filePath))
        # Add File in Folder Root
        folderMgm.AddFileFolderRoot(usr, fObj.id)
        # Update PdfManipulation User Stats
        fMgm.UpdatePdfManipulationCount(usr)

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF Delete Pages error: {0}".format(e))

    return Response(data={"file": fObj, "link": url}, message="OK", statusCode=status.HTTP_200_OK)


@pdfRouter.post("/rotation", response_model=Response, dependencies=[Depends(packageLimit.FileLimitDependency)])
def pdfRotate(
    request: Request,
    query: pdfRotateQueryParams = Depends(),
    file: UploadFile = File(...),
):
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")

    if Validations.isNumber(query.angle) == False:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="angle is failed")

    try:
        usr = request.state.userDict[u'email']
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan
        filePath = ImageHandling.saveImageToTmp(file)

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
            raise Exception("This feature is for premium account. Your current plan is free")
        
        # rotation pdf
        with Pdf.open(filePath, allow_overwriting_input=True) as pdf:
            for page in pdf.pages:
                page.Rotate = int(query.angle)

            pdf.save(filePath)

        # platform Web Add Watermark
        if platform == WEB_PLATFORM_FEATURE.name:
            filePath = PdfHandling().addWatermark(filePath)
        # Upload file in cloud
        fObj = fMgm.CreateFromPath(userId=usr, filePath=str(filePath))
        url = fMgm.GetDownloadLink(fObj=fObj)
        # Clear file 
        ImageHandling().cleanSingleImage(filePathWithName=str(filePath))
        # Add File in Folder Root
        folderMgm.AddFileFolderRoot(usr, fObj.id)
        # Update PdfManipulation User Stats
        fMgm.UpdatePdfManipulationCount(usr)

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF Rotation error: %s" % (e))

    return Response(data={"file": fObj, "link": url}, message="OK", statusCode=status.HTTP_200_OK)


@pdfRouter.post("/merge", response_model=Response, dependencies=[Depends(packageLimit.FileLimitDependency)])
def pdfMerge(
    request: Request,
    files: List[UploadFile] = File(...),
):
    if not all(file.content_type in "application/pdf" for file in files):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")

    try:
        usr = request.state.userDict[u'email']
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
            raise Exception("This feature is for premium account. Your current plan is free")

        strUUID = uuid.uuid4().hex
        outFolder = ImageHandling.printTmpDir()
        arrFileDel = []
        nameFile = f"{strUUID}.pdf"

        # merge pdf
        with Pdf.new() as pdf:
            for file in files:
                filePath = ImageHandling.saveImageToTmp(file)
                arrFileDel.append(filePath)
                with Pdf.open(filePath) as src:
                    pdf.pages.extend(src.pages)

            f = '%s/%s' % (outFolder, nameFile)
            arrFileDel.append(f)
            pdf.save(f)

        # platform Web Add Watermark
        if platform == WEB_PLATFORM_FEATURE.name:
            f = PdfHandling().addWatermark(f)
        # Upload file in cloud
        fObj = fMgm.CreateFromPath(userId=usr, filePath=f)
        url = fMgm.GetDownloadLink(fObj=fObj)
        # Clear file 
        for i in arrFileDel:
            ImageHandling().cleanSingleImage(filePathWithName=str(i))
        # Add File in Folder Root
        folderMgm.AddFileFolderRoot(usr, fObj.id)
        # Update PdfManipulation User Stats
        fMgm.UpdatePdfManipulationCount(usr)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF Merge error: {0}".format(e))

    return Response(data={"file": fObj, "link": url}, message="OK", statusCode=status.HTTP_200_OK)


@pdfRouter.post("/split", response_model=Response, dependencies=[Depends(packageLimit.FileLimitDependency)])
def pdfSplit(
    request: Request,
    file: UploadFile = File(...),
    query: pdfSplitQueryParams = Depends(),
):
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")
    if len(query.ranges) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="ranges is required")

    try:
        usr = request.state.userDict[u'email']
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan
        filePath = ImageHandling.saveImageToTmp(file)

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
            raise Exception("This feature is for premium account. Your current plan is free")
        
        # platform Web Add Watermark
        if platform == WEB_PLATFORM_FEATURE.name:
            filePath = PdfHandling().addWatermark(filePath)

        outFolder = ImageHandling.printTmpDir()
        arrFileDel = [filePath]
        arrFile = []

        # split pdf
        with Pdf.open(filePath) as pdf:
            ranges = query.ranges.split(',')
            for i in ranges:
                with Pdf.new() as dst:
                    pagesArr = i.split('-')
                    strUUID = uuid.uuid4().hex
                    nameFile = '%s/%s.pdf' % (outFolder, strUUID)
                    arrFile.append(nameFile)

                    if len(pagesArr) == 1:
                        dst.pages.append(pdf.pages.p(int(pagesArr[0])))
                        dst.save(nameFile)
                    if len(pagesArr) == 2:
                        for j in range(int(pagesArr[0]), int(pagesArr[1]) + 1):
                            dst.pages.append(pdf.pages.p(j))
                            dst.save(nameFile)

        arrFileDel.extend(arrFile)

        strZipUUID = uuid.uuid4().hex
        pathZip = '%s/%s.zip' % (outFolder, strZipUUID)

        with ZipFile(pathZip, 'w') as zipObj:
            for tf in arrFile:
                zipObj.write(tf)
        arrFileDel.append(pathZip)

        # Upload file in cloud
        fObj = fMgm.CreateFromPath(userId=usr, filePath=pathZip)
        url = fMgm.GetDownloadLink(fObj=fObj)
        # Clear file 
        for i in arrFileDel:
            ImageHandling().cleanSingleImage(filePathWithName=str(i))
        # Add File in Folder Root
        folderMgm.AddFileFolderRoot(usr, fObj.id)
        # Update PdfManipulation User Stats
        fMgm.UpdatePdfManipulationCount(usr)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF Split error: %s" % (e))

    return Response(data={"file": fObj, "link": url}, message="OK", statusCode=status.HTTP_200_OK)


@pdfRouter.post("/sort", response_model=Response, dependencies=[Depends(packageLimit.FileLimitDependency)])
def pdfSort(
    request: Request,
    file: UploadFile = File(...),
    query: pdfSortQueryParams = Depends(),
):
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")
    if len(query.sorts) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="sorts is required")

    try:
        usr = request.state.userDict[u'email']
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan
        filePath = ImageHandling.saveImageToTmp(file)

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
            raise Exception("This feature is for premium account. Your current plan is free")
        
        # Sort pdf
        pagesArr = query.sorts.split(',')
        with Pdf.open(filePath, allow_overwriting_input=True) as pdf:
            with Pdf.new() as dst:
                for page in pagesArr:
                    dst.pages.append(pdf.pages.p(int(page)))
                    dst.save(filePath)

        # platform Web Add Watermark
        if platform == WEB_PLATFORM_FEATURE.name:
            filePath = PdfHandling().addWatermark(filePath)
        # Upload file in cloud
        fObj = fMgm.CreateFromPath(userId=usr, filePath=str(filePath))
        url = fMgm.GetDownloadLink(fObj=fObj)
        # Clean File
        ImageHandling().cleanSingleImage(filePathWithName=str(filePath))
        # Add File in Folder Root
        folderMgm.AddFileFolderRoot(usr, fObj.id)
        # Update PdfManipulation User Stats
        fMgm.UpdatePdfManipulationCount(usr)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF Sort error: %s" % (e))

    return Response(data={"file": fObj, "link": url}, message="OK", statusCode=status.HTTP_200_OK)


@pdfRouter.post("/watermark", response_model=Response, dependencies=[Depends(packageLimit.FileLimitDependency)])
def pdfWatermark(
    request: Request,
    msg: str = Form(...),
    file: UploadFile = File(...),
):
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")
    if len(msg.strip()) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="msg is not None")

    try:
        usr = request.state.userDict[u'email']
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan
        filePath = ImageHandling.saveImageToTmp(file)

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
            raise Exception("This feature is for premium account. Your current plan is free")

        # Watermark pdf
        filePath = PdfHandling().addWatermark(filePath, watermarkText=msg)
        # Upload file in cloud
        fObj = fMgm.CreateFromPath(userId=usr, filePath=str(filePath))
        url = fMgm.GetDownloadLink(fObj=fObj)
        # Clean File
        ImageHandling().cleanSingleImage(filePathWithName=str(filePath))                               
        # Add File in Folder Root
        folderMgm.AddFileFolderRoot(usr, fObj.id)
        # Update PdfManipulation User Stats
        fMgm.UpdatePdfManipulationCount(usr)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF Watermark error: %s" % (e))

    return Response(data={"file": fObj, "link": url}, message="OK", statusCode=status.HTTP_200_OK)


@pdfRouter.post("/add-password", response_model=Response, dependencies=[Depends(packageLimit.FileLimitDependency)])
def pdfProtect(
    request: Request,
    file: UploadFile = File(...),
    password: str = Form(...),
):
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")
    isPassWD = PdfHandling().validationPassword(password)
    if isPassWD == False:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="Password length must be greater than 3 and less than 8, including uppercase, lowercase, special characters and numbers")

    try:
        usr = request.state.userDict[u'email']
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan
        filePath = ImageHandling.saveImageToTmp(file)

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
            raise Exception("This feature is for premium account. Your current plan is free")

        # platform Web Add Watermark
        if platform == WEB_PLATFORM_FEATURE.name:
            filePath = PdfHandling().addWatermark(filePath)

        # Add Password pdf
        with Pdf.open(filePath, allow_overwriting_input=True) as pdf:
            pdf.save(filePath, encryption=Encryption(
                owner=password, user=password, R=6))

        # Upload file in cloud
        fObj = fMgm.CreateFromPath(userId=usr, filePath=str(filePath))
        url = fMgm.GetDownloadLink(fObj=fObj)
        # Clean File
        ImageHandling().cleanSingleImage(filePathWithName=str(filePath))                               
        # Add File in Folder Root
        folderMgm.AddFileFolderRoot(usr, fObj.id)
        # Update PdfManipulation User Stats
        fMgm.UpdatePdfManipulationCount(usr)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF Add Password error: %s" % (e))

    return Response(data={"file": fObj, "link": url}, message="OK", statusCode=status.HTTP_200_OK)


@pdfRouter.post("/reset-password", response_model=Response, dependencies=[Depends(packageLimit.FileLimitDependency)])
def pdfUnlock(
    request: Request,
    file: UploadFile = File(...),
    password: str = Form(...),
):
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")
    isPassWD = PdfHandling().validationPassword(password)
    if isPassWD == False:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid Password")

    try:
        usr = request.state.userDict[u'email']
        filePath = ImageHandling.saveImageToTmp(file)
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
           raise Exception("This feature is for premium account. Your current plan is free")
        
        # reset password pdf
        with Pdf.open(filePath, password=password, allow_overwriting_input=True) as pdf:
            pdf.save(filePath)

        # platform Web Add Watermark
        if platform == WEB_PLATFORM_FEATURE.name:
            filePath = PdfHandling().addWatermark(filePath)

        # Upload file in cloud
        fObj = fMgm.CreateFromPath(userId=usr, filePath=str(filePath))
        url = fMgm.GetDownloadLink(fObj=fObj)

        # Clean File
        ImageHandling().cleanSingleImage(filePathWithName=str(filePath))                               
        # Add File in Folder Root
        folderMgm.AddFileFolderRoot(usr, fObj.id)
        # Update PdfManipulation User Stats
        fMgm.UpdatePdfManipulationCount(usr)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF Reset Password error: %s" % (e))

    return Response(data={"file": fObj, "link": url}, message="OK", statusCode=status.HTTP_200_OK)


@pdfRouter.post("/signature", response_model=Response, dependencies=[Depends(packageLimit.FileLimitDependency)])
async def pdfSignature(
    request: Request,
    query: pdfSignatureQueryParams = Depends(),
    file: UploadFile = File(...),
    signatureImage: UploadFile = File(...),
):
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")
    if signatureImage.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="signatureImage is not type images")
    if Validations.isNumber(query.page) == False:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="page is Failed")
    if Validations.isNumber(query.x) == False:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="x is Failed")
    if Validations.isNumber(query.y) == False:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="y is Failed")

    try:
        usr = request.state.userDict[u'email']
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan
        filePath = ImageHandling.saveImageToTmp(file)

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
            raise Exception("This feature is for premium account. Your current plan is free")

        imagePath = ImageHandling.saveImageToTmp(signatureImage)
        arrFileDel = [filePath, imagePath]
        #  signature pdf
        with Pdf.open(filePath, allow_overwriting_input=True) as pdf:
            destinationPage = Page(pdf.pages.p(int(query.page)))
            w, h = PdfPageSize.get(destinationPage)
            img = ImageHandling.removeBackground(imagePath)
            pdfSignature = PdfHandling().addImageSignature(
                imagePath, (w, h), (float(query.x), float(query.y)))
            arrFileDel.extend([img, pdfSignature])
            with Pdf.open(pdfSignature) as signaturePDF:
                signaturePDF = Page(signaturePDF.pages.p(1))
                destinationPage.add_overlay(signaturePDF)

            pdf.save(filePath)

        # platform Web Add Watermark
        if platform == WEB_PLATFORM_FEATURE.name:
            filePath = PdfHandling().addWatermark(filePath)

        # Upload file in cloud
        fObj = fMgm.CreateFromPath(userId=usr, filePath=str(filePath))
        url = fMgm.GetDownloadLink(fObj=fObj)
        # Clear file 
        for i in arrFileDel:
            ImageHandling().cleanSingleImage(filePathWithName=str(i))
        # Add File in Folder Root
        folderMgm.AddFileFolderRoot(usr, fObj.id)
        # Update PdfManipulation User Stats
        fMgm.UpdatePdfManipulationCount(usr)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF Signature error: %s" % (e))

    return Response(data={"file": fObj, "link": url}, message="OK", statusCode=status.HTTP_200_OK)


@pdfRouter.post("/extract-images", response_model=Response, dependencies=[Depends(packageLimit.FileLimitDependency)])
def pdfExtractImages(
    request: Request,
    file: UploadFile = File(...),
):
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")

    try:
        usr = request.state.userDict[u'email']
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan
        filePath = ImageHandling.saveImageToTmp(file)

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
            raise Exception("This feature is for premium account. Your current plan is free")


        filename = uuid.uuid4().hex
        arrFileDel = [filePath]
        arrFile = []
        outFolder = ImageHandling.printTmpDir()

        # extract images pdf
        with Pdf.open(filePath) as pdf:
            for i, page in enumerate(pdf.pages):
                for j, (name, raw_image) in enumerate(page.images.items()):
                    image = PdfImage(raw_image)
                    out = image.extract_to(
                        fileprefix=f"{outFolder}/{filename}-page{i}-img{j}")
                    arrFile.append(out)

        arrFileDel.extend(arrFile)

        strZipUUID = uuid.uuid4().hex
        pathZip = f"{outFolder}/{strZipUUID}.zip"
        with ZipFile(pathZip, 'w') as zipObj:
            for tf in arrFile:
                zipObj.write(tf)
        arrFileDel.append(pathZip)

        # Upload file in cloud
        fObj = fMgm.CreateFromPath(userId=usr, filePath=pathZip)
        url = fMgm.GetDownloadLink(fObj=fObj)
        # Clear file 
        for i in arrFileDel:
            ImageHandling().cleanSingleImage(filePathWithName=str(i))
        # Add File in Folder Root
        folderMgm.AddFileFolderRoot(usr, fObj.id)
        # Update PdfManipulation User Stats
        fMgm.UpdatePdfManipulationCount(usr)

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF Extract Images error: %s" % (e))

    return Response(data={"file": fObj, "link": url}, message="OK", statusCode=status.HTTP_200_OK)


@pdfRouter.post("/removing-images", response_model=Response, dependencies=[Depends(packageLimit.FileLimitDependency)])
def pdfRemovingImage(
    request: Request,
    file: UploadFile = File(...),
):
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")

    try:
        usr = request.state.userDict[u'email']
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan
        filePath = ImageHandling.saveImageToTmp(file)

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
            raise Exception("This feature is for premium account. Your current plan is free")

        outFolder = ImageHandling.printTmpDir()
        strZipUUID = uuid.uuid4().hex
        pathPdf = f"{outFolder}/{strZipUUID}.pdf"

        # removing-images pdf
        with Pdf.new() as dst:
            with Pdf.open(filePath) as pdf:
                for i, page in enumerate(pdf.pages):
                    for j, (image_name, image) in enumerate(page.images.items()):
                        new_image = pdf.make_stream(b'\xff')
                        new_image.Width, new_image.Height = 1, 1
                        new_image.BitsPerComponent = 1
                        new_image.ImageMask = True
                        new_image.Decode = [0, 1]
                        page.Resources.XObject[image_name] = new_image
                    dst.pages.append(page)
        dst.save(pathPdf)

        # platform Web Add Watermark
        if platform == WEB_PLATFORM_FEATURE.name:
            pathPdf = PdfHandling().addWatermark(pathPdf)

        # Upload file in cloud
        fObj = fMgm.CreateFromPath(userId=usr, filePath=pathPdf)
        url = fMgm.GetDownloadLink(fObj=fObj)

        # Clear file 
        ImageHandling().cleanSingleImage(filePathWithName=str(pathPdf))
        # Add File in Folder Root
        folderMgm.AddFileFolderRoot(usr, fObj.id)
        # Update PdfManipulation User Stats
        fMgm.UpdatePdfManipulationCount(usr)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF Removing Images error: %s" % (e))

    return Response(data={"file": fObj, "link": url}, message="OK", statusCode=status.HTTP_200_OK)


@pdfRouter.post("/docx", response_model=Response)
def pdfToDocx(
    request: Request,
    file: UploadFile = File(...),
):
    if file.content_type != "application/pdf":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")
    try:
        usr = request.state.userDict[u'email']
        filePath = ImageHandling.saveImageToTmp(file)
        tmpFolder = ImageHandling.printTmpDir()
        outFolder = uuid.uuid4().hex
        nameFile = uuid.uuid4().hex
        pathFolder = f"{tmpFolder}/{outFolder}"
        ImageHandling().createSingleDirectory(directoryPath=pathFolder)
        pathXml = f"{pathFolder}/{nameFile}.xml"
        pathDocx = f"{pathFolder}/{nameFile}.docx"

        isErr = PdfaltoHandling().toXML(str(filePath), pathXml)
        if isErr == False:
            raise Exception("Server Error toXML")

        with open(pathXml, 'r') as f:
            xmlData = f.read()

        soup = BeautifulSoup(xmlData, "lxml")

        document = Document()
        section = document.sections[0]

        pages = soup.layout.find_all("page")
        for page in pages:
            pageHeight = float(page.get("height"))
            pageWidth = float(page.get("width"))
            section.page_height = Inches(pageHeight / 72)
            section.page_width = Inches(pageWidth / 72)

            textblocks = page.printspace.find_all("textblock")
            if len(textblocks) > 0:
                for tb in textblocks:
                    ap = document.add_paragraph()
                    ap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    ap.paragraph_format.line_spacing = 1.15

                    tls = tb.find_all('textline')
                    for tl in tls:
                        composedBlockVpos = float(tl.get("vpos"))
                        composedBlockHpos = float(tl.get("hpos"))
                        composedBlockAlignment = -1
                        if (composedBlockAlignment == -1):
                            if (composedBlockHpos > .8 * pageWidth):
                                ap.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                            if (composedBlockHpos > 0 and composedBlockHpos <= .15 * pageWidth):
                                ap.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                            if (composedBlockHpos >= .2 * pageWidth and composedBlockHpos < .8 * pageWidth):
                                ap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        stris = tl.find_all('string')
                        for stri in stris:
                            styleID = stri.get('stylerefs')
                            style = soup.find(id=styleID)
                            t = stri.get('content')
                            run = ap.add_run(f'{t} ')
                            font = run.font
                            font.name = style.get('fontfamily')
                            font.size = Pt(float(style.get('fontsize')))
                            color = style.get('fontcolor')
                            rgb = PdfaltoHandling().hexToRgb(f'#{color}')
                            font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
                            if style.get('fontstyle') == 'bold':
                                font.bold = True
                            if style.get('fontstyle') == 'italics':
                                font.italic = True
                            if style.get('fontstyle') == 'underline':
                                font.underline = True

            document.add_page_break()

        document.save(pathDocx)

        res = fMgm.CreateFromPathAndDownloadLink(userId=usr, filePath=pathDocx)
        ImageHandling().cleanSingleDirectory(directoryPath=pathFolder)
        ImageHandling().cleanSingleImage(filePathWithName=str(filePath))
        fMgm.UpdatePdfManipulationCount(usr)

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="Server Error %s" % (e))

    return Response(data={"link": res}, message="OK", statusCode=status.HTTP_200_OK)


@defers_collector
@pdfRouter.post("/organize", response_model=Response, description="Combination of functions: merge, rorate, sort, delete pages ")
def organizePDFs(
    request: Request,
    files: List[UploadFile] = File(description="List pdf files"),
    task: str = Form(..., description="Defines organize task processing",
                     example="0:1-2#180,1:1-1#90"),
    password: str = Form(default="", description="Create or replace password"),
    filesPassword: str = Form(default="", description="Passwords to access uploaded files",
                              example="0:base64encoded,1:base64encoded"),
):
    res = ""
    arrFileDel = []
    strUUID = uuid.uuid4().hex
    outFolder = ImageHandling.printTmpDir()
    nameFile = f"{strUUID}.pdf"

    defer(lambda: cleanTempFiles(arrFileDel))

    if not all(file.content_type in "application/pdf" for file in files):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="file is not pdf")

    if password != "" and PdfHandling().validationPassword(password) is False:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid Password")

    try:
        usr = request.state.userDict[u'email']
        platform = request.state.platform
        userCurrentPlan = uMgm.Info(usr).current_plan

        if platform == APP_PLATFORM_FEATURE.name and userCurrentPlan == FREE_PLAN_FEATURE.name:
            raise Exception("This feature is for premium account. Your current plan is free")

        filesPw = dict()
        if filesPassword != "":
            filesPw = parseFilesPassword(filesPassword)

        pdfObjs = []
        idx = 0
        for file in files:
            filePath = ImageHandling.saveImageToTmp(file)
            arrFileDel.append(filePath)

            fileStats = os.stat(filePath)
            FeatureLimit.IsValidOrRaiseHttpException(
                usr, "file_capacity", fileStats.st_size)

            fpw = ""
            if filesPw[str(idx)] != "":
                fpw = filesPw[str(idx)]

            pdfObj = Pdf.open(filePath, password=fpw)
            pdfObjs.append(pdfObj)
            idx = idx + 1

        pdfTask = PdfTask(pdfObjs)
        newPdf = pdfTask.buildFromString(task)

        f = '%s/%s' % (outFolder, nameFile)
        arrFileDel.append(f)
        newPdf.save(f)

        if platform == WEB_PLATFORM_FEATURE.name:
            f = PdfHandling().addWatermark(f)

        # Add new password
        if password != "":
            with Pdf.open(f, allow_overwriting_input=True) as pdf:
                pdf.save(filePath, encryption=Encryption(
                    owner=password, user=password, R=6))

        res = fMgm.CreateFromPathAndDownloadLink(userId=usr, filePath=f)
        fMgm.UpdatePdfManipulationCount(usr)

    except TaskError as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Processing pdfs has some errors, details: %s" % (e))
    except ErrorInfoModel as e:
        raise HTTPException(
            status_code=e.http_status, detail=e.message)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="PDF organize error: %s" % (e))

    return Response(data={"link": res}, message="OK", statusCode=status.HTTP_200_OK)


def cleanTempFiles(arrFileDel: List):
    if len(arrFileDel) == 0:
        return

    print("Start clean temp files")
    for i in arrFileDel:
        ImageHandling().cleanSingleImage(str(i))
